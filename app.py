from flask import Flask, render_template, request, redirect, url_for, send_file
from docx import Document
import os

app = Flask(__name__)

# Data list untuk menyimpan nama, alasan, nominal, dan status fix
people_list = []

@app.route('/')
def index():
    return render_template('index.html', people_list=people_list)

@app.route('/add', methods=['POST'])
def add_person():
    name = request.form.get('name')
    reason = request.form.get('reason')
    nominal = request.form.get('nominal')

    # Validasi input nominal
    try:
        nominal = float(nominal) if nominal else 0
    except ValueError:
        nominal = 0

    if name and reason:
        people_list.append({"name": name, "reason": reason, "nominal": nominal, "fix": False})
    return redirect(url_for('index'))

@app.route('/fix/<int:person_id>')
def fix_person(person_id):
    if 0 <= person_id < len(people_list):
        people_list[person_id]['fix'] = True
    return redirect(url_for('index'))

@app.route('/download')
def download():
    # Buat file Word
    document = Document()
    document.add_heading('Daftar Nama Orang', level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header tabel
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nama'
    hdr_cells[2].text = 'Alasan & Nominal'
    hdr_cells[3].text = 'Status'

    # Isi tabel
    for i, person in enumerate(people_list, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = person['name']
        row_cells[2].text = f"{person['reason']} - Rp {person['nominal']:,.2f}"
        row_cells[3].text = 'Fix' if person['fix'] else 'Belum Fix'

    # Simpan file Word
    file_path = "Daftar_Nama_Orang.docx"
    document.save(file_path)

    # Kirim file ke pengguna
    return send_file(file_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
